import re, sys, io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Configure heading styles
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# Add title page
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.space_before = Pt(120)
run = title_para.add_run('OCD CURES RESEARCH')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Comprehensive Treatment & Research Compendium')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

doc.add_page_break()

# Add Table of Contents field
toc_heading = doc.add_heading('Table of Contents', level=1)
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._r.append(fldChar)
run2 = paragraph.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
run2._r.append(instrText)
run3 = paragraph.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
run3._r.append(fldChar2)
run4 = paragraph.add_run('Right-click and select "Update Field" to generate Table of Contents')
run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run4.font.italic = True
run5 = paragraph.add_run()
fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run5._r.append(fldChar3)

doc.add_page_break()

# Read extracted content
with open('_extracted_content.txt', 'r', encoding='utf-8') as f:
    lines = f.readlines()

# Heading detection patterns
category_pattern = re.compile(r'^CATEGORY\s+\d+', re.IGNORECASE)
section_pattern = re.compile(r'^(SECTION\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+)|PART\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+)|TIER\s+\d+|LAYER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+))', re.IGNORECASE)
subcategory_pattern = re.compile(r'^(SUBCATEGORY\s+[A-Z]|CLASS\s+[A-Z]|[A-Z]\.\s+[A-Z]|[IVX]+\.\s+[A-Z])', re.IGNORECASE)
numbered_heading_pattern = re.compile(r'^(\d+\.\s+[A-Z]|[①②③④⑤⑥⑦⑧⑨⑩❶❷❸❹❺❻❼❽❾]\s*[A-Z])')
agent_pattern = re.compile(r'^AGENT\s+\d+:', re.IGNORECASE)
section_letter_pattern = re.compile(r'^[A-Z]\.\s+[A-Z]')
roman_section = re.compile(r'^[IVX]+\.\s+[A-Z]')
dash_section = re.compile(r'^───\s+.+\s+───$')
heading_24_pattern = re.compile(r'^24[A-Z]\s+—')

def classify_paragraph(text, is_bold):
    if not text:
        return 'empty', 0

    # Category = Heading 1
    if category_pattern.match(text):
        return 'heading', 1

    # Major sections
    if section_pattern.match(text):
        return 'heading', 2

    # Sub-sections
    if subcategory_pattern.match(text) and is_bold and len(text) < 120:
        return 'heading', 2

    if agent_pattern.match(text):
        return 'heading', 3

    if heading_24_pattern.match(text) and is_bold:
        return 'heading', 2

    if dash_section.match(text):
        return 'heading', 2

    # Bold uppercase lines that are short = likely headings
    if is_bold and text == text.upper() and len(text) < 100 and len(text) > 3:
        return 'heading', 2

    # Numbered headings like "1. EXPOSURE..." or circled numbers
    if numbered_heading_pattern.match(text) and is_bold and len(text) < 150:
        return 'heading', 3

    # Bold lines starting with letter+dot pattern
    if section_letter_pattern.match(text) and is_bold and len(text) < 120:
        return 'heading', 2

    if roman_section.match(text) and is_bold and len(text) < 120:
        return 'heading', 2

    # Bold short lines could be sub-headings
    if is_bold and len(text) < 80 and len(text) > 3:
        return 'heading', 3

    return 'body', 0

para_count = 0
for line in lines:
    line = line.rstrip('\n')
    if not line or len(line) < 2:
        continue

    marker = line[0]
    text = line[2:]  # skip "B|" or "N|"
    is_bold = (marker == 'B')

    if not text.strip():
        continue

    ptype, level = classify_paragraph(text.strip(), is_bold)

    if ptype == 'heading':
        # Clean up heading text
        heading_text = text.strip()
        # Remove emoji prefixes for cleaner headings
        heading_text = re.sub(r'^[📊⚠️⭐🔵❶❷❸❹❺❻❼❽❾]\s*', '', heading_text)
        doc.add_heading(heading_text, level=level)
    elif ptype == 'body':
        p = doc.add_paragraph(text.strip())
        if is_bold:
            for run in p.runs:
                run.bold = True

    para_count += 1

# Save
output_path = 'OCD_CURES_RESEARCH_Structured.docx'
doc.save(output_path)
print(f'Created {output_path} with {para_count} paragraphs')
